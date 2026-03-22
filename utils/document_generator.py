import os
import json
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BaseGenerator:
    @staticmethod
    def get_filename(name, extension):
        clean_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
        clean_name = clean_name.replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{clean_name}_{timestamp}.{extension}"
        return filename

class PDFGenerator(BaseGenerator):
    @staticmethod
    def generate(resume_data):
        filename = BaseGenerator.get_filename(
            resume_data.get('personal', {}).get('fullName', 'resume'),
            'pdf'
        )
        filepath = os.path.join('/tmp', filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=6,
            alignment=1
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=6,
            topMargin=6,
            bottomMargin=3,
            borderColor=colors.HexColor('#e74c3c'),
            borderWidth=0,
            borderPadding=4
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4
        )
        
        # Personal Information
        personal = resume_data.get('personal', {})
        if personal.get('fullName'):
            story.append(Paragraph(personal['fullName'], title_style))
        
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])
        
        if contact_info:
            story.append(Paragraph(' | '.join(contact_info), normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Summary
        summary = resume_data.get('summary', '')
        if summary:
            story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
            story.append(Paragraph(summary, normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Experience
        experience = resume_data.get('experience', [])
        if experience:
            story.append(Paragraph('EXPERIENCE', heading_style))
            for job in experience:
                if job.get('jobTitle'):
                    job_header = f"<b>{job.get('jobTitle', '')}</b> - {job.get('company', '')}"
                    story.append(Paragraph(job_header, normal_style))
                    job_dates = f"{job.get('startDate', '')} - {job.get('endDate', '')}"
                    story.append(Paragraph(job_dates, normal_style))
                    if job.get('description'):
                        story.append(Paragraph(job['description'], normal_style))
                    story.append(Spacer(1, 0.05*inch))
            story.append(Spacer(1, 0.05*inch))
        
        # Education
        education = resume_data.get('education', [])
        if education:
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in education:
                if edu.get('school'):
                    edu_header = f"<b>{edu.get('school', '')}</b>"
                    story.append(Paragraph(edu_header, normal_style))
                    degree_info = []
                    if edu.get('degree'):
                        degree_info.append(edu['degree'])
                    if edu.get('field'):
                        degree_info.append(edu['field'])
                    if degree_info:
                        story.append(Paragraph(' in '.join(degree_info), normal_style))
                    if edu.get('graduationDate'):
                        story.append(Paragraph(edu['graduationDate'], normal_style))
                    story.append(Spacer(1, 0.05*inch))
            story.append(Spacer(1, 0.05*inch))
        
        # Skills
        skills = resume_data.get('skills', [])
        if skills:
            story.append(Paragraph('SKILLS', heading_style))
            skills_text = ', '.join([s.get('skill', '') for s in skills if s.get('skill')])
            if skills_text:
                story.append(Paragraph(skills_text, normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Certifications
        certifications = resume_data.get('certifications', [])
        if certifications:
            story.append(Paragraph('CERTIFICATIONS', heading_style))
            for cert in certifications:
                if cert.get('name'):
                    story.append(Paragraph(f"<b>{cert['name']}</b>", normal_style))
                    if cert.get('issuer'):
                        story.append(Paragraph(f"Issued by {cert['issuer']}", normal_style))
                    story.append(Spacer(1, 0.05*inch))
        
        # Build PDF
        doc.build(story)
        return filepath

class WordGenerator(BaseGenerator):
    @staticmethod
    def generate(resume_data):
        filename = BaseGenerator.get_filename(
            resume_data.get('personal', {}).get('fullName', 'resume'),
            'docx'
        )
        filepath = os.path.join('/tmp', filename)
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Personal Information
        personal = resume_data.get('personal', {})
        if personal.get('fullName'):
            heading = doc.add_paragraph(personal['fullName'])
            heading.style = 'Heading 1'
            heading_format = heading.paragraph_format
            heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_format.space_after = Pt(2)
        
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(6)
        
        # Summary
        summary = resume_data.get('summary', '')
        if summary:
            doc.add_paragraph('PROFESSIONAL SUMMARY', style='Heading 2')
            doc.add_paragraph(summary)
        
        # Experience
        experience = resume_data.get('experience', [])
        if experience:
            doc.add_paragraph('EXPERIENCE', style='Heading 2')
            for job in experience:
                if job.get('jobTitle'):
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(job.get('jobTitle', ''))
                    job_run.bold = True
                    job_para.add_run(f" - {job.get('company', '')}")
                    
                    doc.add_paragraph(f"{job.get('startDate', '')} - {job.get('endDate', '')}")
                    
                    if job.get('description'):
                        doc.add_paragraph(job['description'])
        
        # Education
        education = resume_data.get('education', [])
        if education:
            doc.add_paragraph('EDUCATION', style='Heading 2')
            for edu in education:
                if edu.get('school'):
                    school_para = doc.add_paragraph()
                    school_run = school_para.add_run(edu.get('school', ''))
                    school_run.bold = True
                    
                    degree_info = []
                    if edu.get('degree'):
                        degree_info.append(edu['degree'])
                    if edu.get('field'):
                        degree_info.append(edu['field'])
                    
                    if degree_info:
                        doc.add_paragraph(' in '.join(degree_info))
                    if edu.get('graduationDate'):
                        doc.add_paragraph(edu['graduationDate'])
        
        # Skills
        skills = resume_data.get('skills', [])
        if skills:
            doc.add_paragraph('SKILLS', style='Heading 2')
            skills_text = ', '.join([s.get('skill', '') for s in skills if s.get('skill')])
            if skills_text:
                doc.add_paragraph(skills_text)
        
        # Certifications
        certifications = resume_data.get('certifications', [])
        if certifications:
            doc.add_paragraph('CERTIFICATIONS', style='Heading 2')
            for cert in certifications:
                if cert.get('name'):
                    cert_para = doc.add_paragraph()
                    cert_run = cert_para.add_run(cert['name'])
                    cert_run.bold = True
                    
                    if cert.get('issuer'):
                        doc.add_paragraph(f"Issued by {cert['issuer']}")
        
        doc.save(filepath)
        return filepath

class HTMLGenerator(BaseGenerator):
    @staticmethod
    def generate(resume_data):
        filename = BaseGenerator.get_filename(
            resume_data.get('personal', {}).get('fullName', 'resume'),
            'html'
        )
        filepath = os.path.join('/tmp', filename)
        
        html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
        }
        
        .container {
            max-width: 8.5in;
            height: 11in;
            margin: 20px auto;
            padding: 40px;
            background-color: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        
        .header {
            text-align: center;
            margin-bottom: 20px;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 10px;
        }
        
        .name {
            font-size: 28px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 5px;
        }
        
        .contact-info {
            font-size: 11px;
            color: #666;
        }
        
        .section {
            margin-bottom: 15px;
        }
        
        .section-title {
            font-size: 13px;
            font-weight: bold;
            color: #2c3e50;
            border-bottom: 2px solid #e74c3c;
            padding-bottom: 5px;
            margin-bottom: 10px;
        }
        
        .job, .education-item, .cert {
            margin-bottom: 12px;
        }
        
        .job-title {
            font-weight: bold;
            color: #34495e;
        }
        
        .job-company {
            color: #666;
            font-style: italic;
        }
        
        .job-dates {
            color: #999;
            font-size: 10px;
        }
        
        .job-description {
            color: #555;
            font-size: 11px;
            margin-top: 5px;
        }
        
        .skills {
            font-size: 11px;
            line-height: 1.8;
        }
        
        @media print {
            body {
                margin: 0;
                padding: 0;
            }
            .container {
                margin: 0;
                box-shadow: none;
                height: auto;
            }
        }
    </style>
</head>
<body>
    <div class="container">
"""
        
        # Personal Information
        personal = resume_data.get('personal', {})
        if personal.get('fullName') or personal.get('email'):
            html_content += '        <div class="header">\n'
            if personal.get('fullName'):
                html_content += f'            <div class="name">{personal["fullName"]}</div>\n'
            
            contact_info = []
            if personal.get('email'):
                contact_info.append(personal['email'])
            if personal.get('phone'):
                contact_info.append(personal['phone'])
            if personal.get('location'):
                contact_info.append(personal['location'])
            
            if contact_info:
                html_content += f'            <div class="contact-info">{" | ".join(contact_info)}</div>\n'
            
            html_content += '        </div>\n'
        
        # Summary
        summary = resume_data.get('summary', '')
        if summary:
            html_content += '        <div class="section">\n'
            html_content += '            <div class="section-title">PROFESSIONAL SUMMARY</div>\n'
            html_content += f'            <p style="font-size: 11px;">{summary}</p>\n'
            html_content += '        </div>\n'
        
        # Experience
        experience = resume_data.get('experience', [])
        if experience:
            html_content += '        <div class="section">\n'
            html_content += '            <div class="section-title">EXPERIENCE</div>\n'
            for job in experience:
                if job.get('jobTitle'):
                    html_content += '            <div class="job">\n'
                    html_content += f'                <div class="job-title">{job.get("jobTitle", "")}</div>\n'
                    html_content += f'                <div class="job-company">{job.get("company", "")}</div>\n'
                    html_content += f'                <div class="job-dates">{job.get("startDate", "")} - {job.get("endDate", "")}</div>\n'
                    if job.get('description'):
                        html_content += f'                <div class="job-description">{job["description"]}</div>\n'
                    html_content += '            </div>\n'
            html_content += '        </div>\n'
        
        # Education
        education = resume_data.get('education', [])
        if education:
            html_content += '        <div class="section">\n'
            html_content += '            <div class="section-title">EDUCATION</div>\n'
            for edu in education:
                if edu.get('school'):
                    html_content += '            <div class="education-item">\n'
                    html_content += f'                <div style="font-weight: bold;">{edu.get("school", "")}</div>\n'
                    
                    degree_info = []
                    if edu.get('degree'):
                        degree_info.append(edu['degree'])
                    if edu.get('field'):
                        degree_info.append(edu['field'])
                    
                    if degree_info:
                        html_content += f'                <div style="font-size: 11px;">{" in ".join(degree_info)}</div>\n'
                    if edu.get('graduationDate'):
                        html_content += f'                <div style="font-size: 10px; color: #999;">{edu["graduationDate"]}</div>\n'
                    html_content += '            </div>\n'
            html_content += '        </div>\n'
        
        # Skills
        skills = resume_data.get('skills', [])
        if skills:
            skills_list = [s.get('skill', '') for s in skills if s.get('skill')]
            if skills_list:
                html_content += '        <div class="section">\n'
                html_content += '            <div class="section-title">SKILLS</div>\n'
                html_content += f'            <div class="skills">{", ".join(skills_list)}</div>\n'
                html_content += '        </div>\n'
        
        # Certifications
        certifications = resume_data.get('certifications', [])
        if certifications:
            html_content += '        <div class="section">\n'
            html_content += '            <div class="section-title">CERTIFICATIONS</div>\n'
            for cert in certifications:
                if cert.get('name'):
                    html_content += '            <div class="cert">\n'
                    html_content += f'                <div style="font-weight: bold; font-size: 11px;">{cert["name"]}</div>\n'
                    if cert.get('issuer'):
                        html_content += f'                <div style="font-size: 10px; color: #666;">Issued by {cert["issuer"]}</div>\n'
                    html_content += '            </div>\n'
            html_content += '        </div>\n'
        
        html_content += """    </div>
</body>
</html>"""
        
        with open(filepath, 'w') as f:
            f.write(html_content)
        
        return filepath
